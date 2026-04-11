"""
Generate comprehensive project documentation in Word format.
Run: python scripts/generate_docs.py
"""
import os
import sys
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex: str):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elem)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '0078D4')

    # Data rows
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, 'F0F4F8')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_building_block_diagram(doc):
    """Add a text-based building block architecture diagram."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Use a single-cell table with monospaced font for the diagram
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'F8FAFC')

    diagram_text = """
┌─────────────────────────────────────────────────────────────────────────────────┐
│                              EFESTO — AI FABRYC                                  │
│                              Building Block View                                 │
└─────────────────────────────────────────────────────────────────────────────────┘

  ┌──────────────────────────────────────────────────────────────────────────────┐
  │  PRESENTATION LAYER                                                          │
  │  ┌─────────────────────────────────────────────────────────────────────────┐ │
  │  │  React SPA (Vite + Tailwind CSS)                                       │ │
  │  │  ┌────────────┐ ┌──────────────┐ ┌──────────┐ ┌─────────┐ ┌─────────┐ │ │
  │  │  │  Dashboard  │ │ New Analysis │ │  Report  │ │ Quality │ │ History │ │ │
  │  │  │   (Home)    │ │  (Upload +   │ │ (Results │ │  (Sonar │ │ (Past   │ │ │
  │  │  │             │ │  Templates)  │ │  View)   │ │  Gate)  │ │ Runs)   │ │ │
  │  │  └────────────┘ └──────────────┘ └──────────┘ └─────────┘ └─────────┘ │ │
  │  └────────────────────────────────────┬──────────────────────────────────┘ │
  └───────────────────────────────────────┼──────────────────────────────────────┘
                                          │  REST API (JSON)
                                          ▼
  ┌──────────────────────────────────────────────────────────────────────────────┐
  │  API LAYER (FastAPI)                                                         │
  │  ┌──────────────────┐  ┌────────────────────┐  ┌──────────────────────────┐ │
  │  │ POST /start      │  │ POST /quick-scan   │  │ GET /{session}/status    │ │
  │  │ (async analysis) │  │ (sync < 120s)      │  │ GET /{session} (report)  │ │
  │  └────────┬─────────┘  └────────┬───────────┘  └──────────────────────────┘ │
  └───────────┼──────────────────────┼──────────────────────────────────────────┘
              │                      │
              ▼                      ▼
  ┌──────────────────────────────────────────────────────────────────────────────┐
  │  ORCHESTRATOR AGENT                                                          │
  │                                                                              │
  │  Phase 1 (Sequential)          Phase 2 (Parallel, max 4 concurrent)          │
  │  ┌────────────────────────┐    ┌────────────────────────────────────────────┐│
  │  │ 1. Code Analyzer       │───▶│ 3. Cost Optimizer    5. GAP Analyzer      ││
  │  │ 2. Infra Analyzer      │    │ 4. Migration Planner 6. WAF Reviewer      ││
  │  └────────────────────────┘    │                      7. Quality Analyzer   ││
  │                                └────────────────────────────────────────────┘│
  │                                                                              │
  │  Phase 3: GPT-4o Synthesis ──▶ Final Report (JSON)                          │
  └──────────────────────────────────────────────────────────────────────────────┘
              │                                        │
              ▼                                        ▼
  ┌──────────────────────┐              ┌──────────────────────────────────────┐
  │  EXECUTION MODES     │              │  TOOLS & KNOWLEDGE                   │
  │  ┌────────────────┐  │              │  ┌──────────────┐ ┌───────────────┐  │
  │  │ Foundry Mode   │  │              │  │ Code Scanner │ │ Pricing Calc  │  │
  │  │ (AI Agent Svc  │  │              │  │ (static      │ │ (Azure Retail │  │
  │  │  persistence,  │  │              │  │  analysis)   │ │  Prices API)  │  │
  │  │  threading,    │  │              │  └──────────────┘ └───────────────┘  │
  │  │  file search)  │  │              │  ┌──────────────┐ ┌───────────────┐  │
  │  ├────────────────┤  │              │  │ IaC Parser   │ │ AI Search     │  │
  │  │ Direct Mode    │  │              │  │ (Terraform,  │ │ (CAF, WAF,    │  │
  │  │ (Azure OpenAI  │  │              │  │  Bicep, ARM, │ │  migration    │  │
  │  │  lower latency │  │              │  │  K8s, Docker)│ │  patterns)    │  │
  │  │  completions)  │  │              │  └──────────────┘ └───────────────┘  │
  │  └────────────────┘  │              └──────────────────────────────────────┘
  └──────────────────────┘
              │
              ▼
  ┌──────────────────────────────────────────────────────────────────────────────┐
  │  DATA & INFRASTRUCTURE LAYER (Azure)                                         │
  │  ┌──────────────┐ ┌──────────────┐ ┌──────────────┐ ┌────────────────────┐  │
  │  │ MongoDB 7    │ │ Azure OpenAI │ │ Azure AI     │ │ Azure AI Foundry   │  │
  │  │ (sessions &  │ │ (GPT-4o +    │ │ Search       │ │ (Hub + Project     │  │
  │  │  reports)    │ │  GPT-4o-mini)│ │ (knowledge)  │ │  + Connections)    │  │
  │  └──────────────┘ └──────────────┘ └──────────────┘ └────────────────────┘  │
  │  ┌──────────────┐ ┌──────────────┐ ┌──────────────┐ ┌────────────────────┐  │
  │  │ Container    │ │ Key Vault    │ │ VNet + NSGs  │ │ Log Analytics +    │  │
  │  │ Apps (API +  │ │ (secrets,    │ │ (Private     │ │ App Insights       │  │
  │  │  auto-scale) │ │  Managed ID) │ │  Endpoints)  │ │ (monitoring)       │  │
  │  └──────────────┘ └──────────────┘ └──────────────┘ └────────────────────┘  │
  └──────────────────────────────────────────────────────────────────────────────┘
"""

    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(diagram_text.strip())
    run.font.name = 'Consolas'
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


# ── Main document generation ─────────────────────────────────────────────────

def generate_documentation():
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Styles ────────────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.color.rgb = RGBColor(0x00, 0x30, 0x67)

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Efesto\nAI Fabryc')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Documentazione Tecnica di Architettura')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x00, 0x30, 0x67)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Versione 1.0.0 — Marzo 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run('Powered by Azure AI Foundry · Follows CAF & WAF')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)
    run.italic = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('Indice', level=1)
    toc_items = [
        ('1.', 'Executive Summary'),
        ('2.', 'Scopo del Progetto'),
        ('3.', 'Architettura della Soluzione'),
        ('3.1', 'Diagramma Building-Block'),
        ('3.2', 'Presentation Layer (Frontend)'),
        ('3.3', 'API Layer (Backend)'),
        ('3.4', 'Orchestrator e Multi-Agent Framework'),
        ('3.5', 'Specialist Agents'),
        ('3.6', 'Tools e Knowledge Base'),
        ('3.7', 'Data e Infrastructure Layer'),
        ('4.', 'Modalità di Esecuzione'),
        ('5.', 'Quality Analyzer e SonarQube Integration'),
        ('6.', 'Flusso di Esecuzione End-to-End'),
        ('7.', 'Infrastructure as Code (Bicep)'),
        ('8.', 'Sicurezza e Compliance'),
        ('9.', 'DevOps e Deployment'),
        ('10.', 'Scenari d\'Uso Supportati'),
        ('11.', 'Stack Tecnologico'),
    ]
    for num, title_text in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}  {title_text}')
        run.font.size = Pt(11)
        if not '.' in num or num.endswith('.'):
            run.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('1. Executive Summary', level=1)

    doc.add_paragraph(
        'Efesto AI Fabryc è un sistema multi-agente basato su intelligenza '
        'artificiale progettato per fornire analisi architetturali complete, pianificazione di '
        'migrazione cloud, ottimizzazione dei costi e revisione del Well-Architected Framework (WAF) '
        'di Azure. Il sistema utilizza 7 agenti specializzati coordinati da un orchestratore '
        'che esegue analisi in fasi sequenziali e parallele, producendo report dettagliati con '
        'raccomandazioni azionabili, roadmap di migrazione e stime di risparmio.'
    )

    doc.add_paragraph(
        'La piattaforma è costruita su Azure AI Foundry e supporta due modalità di esecuzione: '
        'Foundry Mode (con persistenza e threading tramite Azure AI Agent Service) e Direct Mode '
        '(chiamate dirette ad Azure OpenAI per latenza ridotta). L\'architettura segue i principi '
        'del Cloud Adoption Framework (CAF) e del Well-Architected Framework (WAF) di Azure, '
        'con sicurezza zero-secret basata su Managed Identity e Private Endpoints.'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 2. SCOPO DEL PROGETTO
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('2. Scopo del Progetto', level=1)

    doc.add_paragraph(
        'Il progetto nasce dall\'esigenza di automatizzare e standardizzare il processo di '
        'analisi architetturale cloud, tipicamente svolto manualmente da team di Solution Architect '
        'in settimane di lavoro. Efesto AI Fabryc riduce questo processo a minuti, '
        'fornendo un\'analisi multi-dimensionale coerente e ripetibile.'
    )

    doc.add_heading('Obiettivi principali', level=2)
    objectives = [
        'Analisi automatica del codice sorgente per identificare linguaggi, framework, accoppiamento cloud e debito tecnico',
        'Parsing dell\'Infrastructure as Code (Terraform, Bicep, ARM, K8s) per mappare risorse e valutare la postura di sicurezza',
        'Ottimizzazione dei costi con analisi FinOps: right-sizing, reserved instances, risparmio PaaS',
        'Pianificazione della migrazione secondo la metodologia CAF 6Rs con wave planning e risk register',
        'Analisi GAP tra stato attuale e stato target su 7 dimensioni',
        'Revisione WAF su tutti i 5 pilastri con scoring e raccomandazioni',
        'Analisi della qualità del codice e IaC con regole SonarQube-compatibili (Quality Analyzer)',
        'Generazione di report esecutivi con roadmap, azioni prioritizzate e stime di risparmio',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_heading('Target audience', level=2)
    doc.add_paragraph(
        'Il sistema è progettato per CTO, VP of Engineering, Cloud Architects, e team di migrazione '
        'che necessitano di valutazioni rapide e comprehensive per prendere decisioni informate '
        'su migrazioni cloud, modernizzazione applicativa e ottimizzazione infrastrutturale.'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 3. ARCHITETTURA DELLA SOLUZIONE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('3. Architettura della Soluzione', level=1)

    doc.add_paragraph(
        'L\'architettura segue un pattern a 4 livelli (Presentation, API, Agent Orchestration, '
        'Data & Infrastructure) con separazione netta delle responsabilità. Il frontend React SPA '
        'comunica con il backend FastAPI tramite REST API JSON. Il backend orchestra 7 agenti AI '
        'specializzati che utilizzano Azure OpenAI GPT-4o per l\'analisi e la sintesi dei risultati.'
    )

    # 3.1 Building Block Diagram
    doc.add_heading('3.1 Diagramma Building-Block', level=2)
    doc.add_paragraph(
        'Il seguente diagramma mostra la struttura a blocchi dell\'intera architettura, '
        'dalle pagine frontend fino ai servizi Azure sottostanti:'
    )
    add_building_block_diagram(doc)

    doc.add_paragraph()

    # 3.2 Presentation Layer
    doc.add_heading('3.2 Presentation Layer (Frontend)', level=2)
    doc.add_paragraph(
        'Il frontend è una Single Page Application (SPA) costruita con React 18, Vite come bundler, '
        'e Tailwind CSS per lo styling. Utilizza un design system dark-mode coerente con componenti '
        'da Radix UI e icone Lucide.'
    )

    add_styled_table(doc,
        ['Pagina', 'Route', 'Funzionalità'],
        [
            ['Dashboard (Home)', '/', 'Overview del sistema, cards agenti, use cases, link rapidi ai report demo'],
            ['New Analysis', '/analysis', 'Form multi-step: info progetto, selezione scope, upload artifacts (drag & drop), template demo precompilati'],
            ['Report', '/report/:id', 'Visualizzazione risultati: maturity score, executive summary, strategy, savings, findings, risks, azioni, roadmap'],
            ['Architecture', '/architecture', 'Diagramma SVG interattivo dell\'architettura con highlights su execution modes, orchestration e security'],
            ['Code Quality', '/quality', 'Dashboard SonarQube-style con overview metriche, lista issues, coverage per modulo, quality gate'],
            ['History', '/history', 'Lista sessioni di analisi passate con status badge e link ai report'],
        ],
        col_widths=[3.5, 3, 10]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Tecnologie chiave: React Router per il routing SPA, TanStack Query per il data fetching '
        'con polling automatico, Axios per le chiamate HTTP, react-dropzone per l\'upload file, '
        'e react-hot-toast per le notifiche. Il frontend supporta una modalità Demo che funziona '
        'senza backend tramite mock API con dati realistici precaricati.'
    )

    # 3.3 API Layer
    doc.add_heading('3.3 API Layer (Backend)', level=2)
    doc.add_paragraph(
        'Il backend è costruito con FastAPI (Python 3.11) e espone una REST API asincrona. '
        'Utilizza Pydantic per la validazione dei modelli, structlog per il logging strutturato, '
        'e Motor (driver asincrono MongoDB) per la persistenza.'
    )

    add_styled_table(doc,
        ['Endpoint', 'Metodo', 'Descrizione'],
        [
            ['/api/analysis/start', 'POST', 'Avvia analisi asincrona. Restituisce session_id immediatamente; l\'analisi viene eseguita in background.'],
            ['/api/analysis/quick-scan', 'POST', 'Analisi sincrona (timeout 120s). Per progetti piccoli (< 10 file).'],
            ['/api/analysis/{session_id}', 'GET', 'Recupera il report completo di una sessione completata.'],
            ['/api/analysis/{session_id}/status', 'GET', 'Stato corrente della sessione (running/completed/failed).'],
            ['/api/analysis/', 'GET', 'Lista tutte le sessioni di analisi.'],
            ['/health', 'GET', 'Health check con lista agenti disponibili e stato connessione Foundry.'],
        ],
        col_widths=[5, 2, 9.5]
    )

    # 3.4 Orchestrator
    doc.add_heading('3.4 Orchestrator e Multi-Agent Framework', level=2)
    doc.add_paragraph(
        'L\'Orchestrator è il componente centrale che coordina tutti gli agenti specializzati. '
        'Implementa una strategia di esecuzione a 3 fasi:'
    )

    p = doc.add_paragraph()
    run = p.add_run('Phase 1 — Analisi Sequenziale (Code Analyzer → Infra Analyzer)')
    run.bold = True
    doc.add_paragraph(
        'I primi due agenti vengono eseguiti in sequenza perché producono il contesto necessario '
        'per tutti gli agenti successivi. Il Code Analyzer identifica linguaggi, framework e '
        'accoppiamento cloud; l\'Infra Analyzer mappa le risorse IaC e la postura di sicurezza.',
    )

    p = doc.add_paragraph()
    run = p.add_run('Phase 2 — Analisi Parallela (5 agenti concorrenti)')
    run.bold = True
    doc.add_paragraph(
        'I 5 agenti rimanenti (Cost Optimizer, Migration Planner, GAP Analyzer, WAF Reviewer, '
        'Quality Analyzer) vengono eseguiti in parallelo con un semaforo che limita la concorrenza '
        'al valore configurabile agent_parallel_limit (default: 4). Ciascun agente riceve i '
        'risultati della Phase 1 come contesto aggiuntivo.',
    )

    p = doc.add_paragraph()
    run = p.add_run('Phase 3 — Sintesi del Report')
    run.bold = True
    doc.add_paragraph(
        'L\'Orchestrator aggrega tutti i risultati e li invia a GPT-4o con un prompt di sintesi '
        'che produce il report finale in formato JSON strutturato: executive summary, maturity score, '
        'key findings, critical risks, recommended strategy, top 10 actions, e roadmap a fasi.',
    )

    # 3.5 Specialist Agents
    doc.add_heading('3.5 Specialist Agents', level=2)
    doc.add_paragraph(
        'Ogni agente specializzato estende la classe astratta BaseAgent che fornisce: '
        'caricamento del prompt YAML, gestione credenziali (Managed Identity o API Key), '
        'client Azure OpenAI/Foundry, e pattern run/parse con timing e error handling.'
    )

    add_styled_table(doc,
        ['Agente', 'Colore', 'Responsabilità', 'Output'],
        [
            ['Code Analyzer', 'Blu', 'Analisi codice sorgente: linguaggi, framework, SDK cloud, pattern architetturali, debito tecnico, readiness containerizzazione', 'Technology inventory, coupling score, 12-factor compliance, migration impact'],
            ['Infra Analyzer', 'Verde', 'Parsing IaC (Terraform, Bicep, ARM, K8s, Docker Compose): inventario risorse, networking, security posture, compute, data layer', 'Resource map, service mapping Azure, networking topology, cost indicators'],
            ['Cost Optimizer', 'Ambra', 'Analisi FinOps: costi correnti vs target, opportunità right-sizing, reserved instances, Azure Hybrid Benefit, ROI migrazione', 'Current/target cost comparison, optimization opportunities, savings estimate'],
            ['Migration Planner', 'Viola', 'Strategia CAF 6Rs, wave planning, landing zone design, risk register, tooling recommendations', '6Rs classification, migration waves, risk register, landing zone blueprint'],
            ['GAP Analyzer', 'Arancione', 'Analisi gap stato attuale vs target su 7 dimensioni: funzionale, operativa, sicurezza, performance, DevX, costo, organizzativa', 'Gap matrix, remediation priorities, effort estimates per dimension'],
            ['WAF Reviewer', 'Rosa', 'Scoring su tutti i 5 pilastri WAF: Reliability, Security, Cost Optimization, Operational Excellence, Performance Efficiency', 'Pillar scores (1-5), findings con severity, recommendations con Azure service references'],
            ['Quality Analyzer', 'Cyan', 'Analisi statica SonarQube-level su codice E IaC: bugs, vulnerabilities, code smells, security hotspots, 12 regole IaC-specifiche', 'Quality gate (PASS/FAIL), issues list, coverage estimate, technical debt in ore'],
        ],
        col_widths=[2.5, 1.5, 6, 6.5]
    )

    # 3.6 Tools & Knowledge
    doc.add_heading('3.6 Tools e Knowledge Base', level=2)

    add_styled_table(doc,
        ['Tool', 'Tipo', 'Descrizione'],
        [
            ['Code Scanner', 'Static Analysis', 'Rileva linguaggi per estensione, framework (React, Vue, Angular, FastAPI, Django, Spring, ecc.), SDK cloud (AWS, Azure, GCP), compliance 12-factor, scansione secrets hardcoded'],
            ['IaC Parser', 'Infrastructure', 'Parsing multi-formato: Terraform HCL, Bicep, ARM JSON, Kubernetes YAML, Docker Compose, CloudFormation. Estrae risorse e le mappa a 49+ tipi Azure/AWS/GCP'],
            ['Pricing Calculator', 'FinOps', 'Client per Azure Retail Prices API (https://prices.azure.com). Fallback estimates per 20+ servizi Azure. Calcoli Reserved Instance e VM-specific costs'],
            ['Azure AI Search', 'Knowledge Base', 'Indici semantici per CAF guidelines, WAF pillars, e migration patterns. Utilizzato come RAG (Retrieval Augmented Generation) dagli agenti tramite FileSearchTool'],
        ],
        col_widths=[3, 2.5, 11]
    )

    # 3.7 Data & Infrastructure
    doc.add_heading('3.7 Data e Infrastructure Layer', level=2)
    doc.add_paragraph(
        'Il layer di persistenza e infrastruttura è interamente su Azure, gestito tramite Bicep IaC:'
    )

    add_styled_table(doc,
        ['Servizio', 'Ruolo', 'Note'],
        [
            ['MongoDB 7', 'Sessioni e report', 'Self-hosted su Container Apps con Azure Files per la persistenza. Scelta vs Cosmos DB per costo e portabilità'],
            ['Azure OpenAI', 'LLM backbone', 'Deployment GPT-4o (analisi principale) e GPT-4o-mini (task leggeri). API version 2025-01-01-preview'],
            ['Azure AI Search', 'Knowledge base', '3 indici: caf-guidelines, waf-pillars, migration-patterns. Ricerca semantica per RAG'],
            ['Azure AI Foundry', 'Agent orchestration', 'Hub + Project con connessioni a OpenAI e AI Search. Supporta persistence e file search'],
            ['Container Apps', 'Compute', 'FastAPI API + Client Nginx. Auto-scale 0-10 repliche. Managed Identity per auth'],
            ['Key Vault', 'Secret management', 'Zero-secret architecture: tutti i secrets in KV, accesso via Managed Identity RBAC'],
            ['VNet + NSGs', 'Networking', 'Hub-spoke topology con Private Endpoints per tutti i servizi PaaS. No public internet exposure'],
            ['Log Analytics + App Insights', 'Monitoring', 'Diagnostics con 90 giorni retention. Application performance monitoring'],
        ],
        col_widths=[3.5, 3, 10]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. MODALITÀ DI ESECUZIONE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('4. Modalità di Esecuzione', level=1)

    doc.add_paragraph(
        'Il framework supporta due modalità di esecuzione mutuamente esclusive, selezionabili '
        'per sessione tramite il flag use_foundry_mode nella request:'
    )

    add_styled_table(doc,
        ['Caratteristica', 'Foundry Mode', 'Direct Mode'],
        [
            ['Servizio', 'Azure AI Foundry Agent Service', 'Azure OpenAI Chat Completions'],
            ['Persistenza', 'Thread e messaggi persistiti lato server', 'Nessuna (stateless)'],
            ['File Search', 'FileSearchTool integrato (search su documenti caricati)', 'Non disponibile'],
            ['Latenza', 'Maggiore (creazione agent + thread per ogni run)', 'Inferiore (singola chiamata API)'],
            ['Costo', 'Superiore (agent hosting + storage)', 'Inferiore (solo token consumati)'],
            ['Use case ideale', 'Analisi complesse con molti file e necessità di audit trail', 'Quick scan, demo, sviluppo locale'],
            ['Cleanup', 'Agent eliminato dopo ogni run (delete_agent)', 'Nessun cleanup necessario'],
        ],
        col_widths=[3.5, 6.5, 6.5]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 5. QUALITY ANALYZER
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('5. Quality Analyzer e SonarQube Integration', level=1)

    doc.add_paragraph(
        'Il Quality Analyzer è il 7° agente specializzato, unico nel suo genere perché analizza '
        'sia il codice applicativo sia l\'Infrastructure as Code con un set di regole '
        'SonarQube-compatibili. Viene eseguito in Phase 2 (parallelo) in ogni scenario di analisi.'
    )

    doc.add_heading('Regole IaC-Specifiche', level=2)
    add_styled_table(doc,
        ['Rule ID', 'Descrizione', 'Severity'],
        [
            ['iac:S1', 'Resource mancante dei tag obbligatori (Environment, Application, CostCenter)', 'MAJOR'],
            ['iac:S2', 'Valori hardcoded che dovrebbero essere variabili/parametri', 'MINOR'],
            ['iac:S3', 'Ruoli IAM/RBAC eccessivamente permissivi (es. Owner/Contributor su subscription)', 'CRITICAL'],
            ['iac:S4', 'Endpoint pubblici senza WAF o protezione NSG', 'CRITICAL'],
            ['iac:S5', 'Storage senza encryption at rest', 'MAJOR'],
            ['iac:S6', 'Logging diagnostico/audit mancante', 'MAJOR'],
            ['iac:S7', 'Nessuna segmentazione di rete (flat network)', 'MAJOR'],
            ['iac:S8', 'Secrets in chiaro (connection strings, password, API keys)', 'BLOCKER'],
            ['iac:S9', 'Health probes mancanti su risorse compute', 'MINOR'],
            ['iac:S10', 'Nessuna configurazione auto-scaling', 'MINOR'],
            ['iac:S11', 'Policy di backup/retention mancanti', 'MAJOR'],
            ['iac:S12', 'API versions o SKU deprecati', 'MINOR'],
        ],
        col_widths=[2, 11, 3.5]
    )

    doc.add_paragraph()
    doc.add_heading('Quality Gate', level=2)
    doc.add_paragraph(
        'Il Quality Gate applica le seguenti condizioni per determinare PASSED/FAILED:'
    )
    conditions = [
        'Nessun bug di severity BLOCKER o CRITICAL',
        'Nessuna vulnerability CRITICAL',
        'Coverage stimata ≥ 60%',
        'Duplicazione codice < 5%',
        'Rating di maintainability ≥ B',
    ]
    for cond in conditions:
        doc.add_paragraph(cond, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        'Il progetto include anche un file sonar-project.properties per l\'integrazione con '
        'un server SonarQube reale, configurato per scansionare sia src/ (Python) che client/src/ (TypeScript).'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 6. FLUSSO END-TO-END
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('6. Flusso di Esecuzione End-to-End', level=1)

    steps = [
        ('1. Upload Artifacts', 'L\'utente accede alla pagina New Analysis, seleziona un template demo o carica manualmente file di codice (.py, .js, .ts, .java, .cs, .go) e IaC (.tf, .bicep, .yaml, .json). Compila le informazioni del progetto: nome, source cloud, costo mensile attuale, contesto aggiuntivo.'),
        ('2. Selezione Scope', 'L\'utente sceglie "Full Analysis" (tutti i 7 agenti) oppure seleziona specifici agenti. Opzionalmente attiva il Foundry Mode per persistenza completa.'),
        ('3. Avvio Analisi', 'La UI invia una POST a /api/analysis/start (asincrono) o /quick-scan (sincrono). Il backend crea una sessione in MongoDB e avvia l\'orchestrator in background.'),
        ('4. Phase 1 — Analisi Sequenziale', 'L\'Orchestrator esegue Code Analyzer e Infra Analyzer in sequenza. I risultati vengono iniettati nel contesto condiviso per arricchire gli agenti della Phase 2.'),
        ('5. Phase 2 — Analisi Parallela', '5 agenti (Cost, Migration, GAP, WAF, Quality) vengono lanciati in parallelo con semaphore. Ogni agente riceve il contesto completo inclusi i risultati Phase 1.'),
        ('6. Phase 3 — Sintesi', 'L\'Orchestrator aggrega i 7 risultati e li invia a GPT-4o con un prompt di sintesi strutturato. Il modello produce il report finale JSON con executive summary, actions e roadmap.'),
        ('7. Report Storage', 'Il report viene salvato in MongoDB (collection reports) e lo stato della sessione aggiornato a "completed".'),
        ('8. Visualizzazione', 'Il frontend effettua polling su /api/analysis/{sessionId} ogni 3 secondi. Alla ricezione del report, visualizza tutte le sezioni: score, strategy, savings, findings, risks, actions, roadmap.'),
    ]
    for title_text, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(title_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)
        doc.add_paragraph(desc)

    # ══════════════════════════════════════════════════════════════════════════
    # 7. INFRASTRUCTURE AS CODE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('7. Infrastructure as Code (Bicep)', level=1)

    doc.add_paragraph(
        'L\'infrastruttura Azure è definita interamente in Bicep con deployment a scope subscription. '
        'Il file main.bicep orchestra 8 moduli nested in ordine di dipendenza:'
    )

    add_styled_table(doc,
        ['Modulo', 'File', 'Risorse Create'],
        [
            ['1. Networking', 'networking.bicep', 'VNet, Subnet (api, data, private-endpoints), NSGs, regole di sicurezza'],
            ['2. Key Vault', 'keyvault.bicep', 'Key Vault con Private Endpoint, RBAC policies, diagnostics'],
            ['3. Monitoring', 'monitoring.bicep', 'Log Analytics Workspace, Application Insights'],
            ['4. OpenAI', 'openai.bicep', 'Azure OpenAI con deployment GPT-4o e GPT-4o-mini'],
            ['5. AI Search', 'ai-search.bicep', 'Azure AI Search (knowledge base per CAF/WAF/patterns)'],
            ['6. Container Apps Env', 'container-apps-env.bicep', 'Managed Environment per Container Apps'],
            ['7. MongoDB', 'mongodb.bicep', 'MongoDB 7 su Container Apps con Azure Files persistence'],
            ['8. AI Foundry', 'ai-foundry.bicep', 'AI Hub + Project con connessioni a OpenAI e AI Search'],
            ['9. Container Apps', 'container-apps.bicep', 'FastAPI API container (scale 0→10, Managed Identity)'],
        ],
        col_widths=[3, 3.5, 10]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Sono disponibili parametri per ambienti dev e prod (infra/parameters/dev.bicepparam, prod.bicepparam) '
        'con differenze su scaling, ridondanza e networking.'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 8. SICUREZZA E COMPLIANCE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('8. Sicurezza e Compliance', level=1)

    security_items = [
        ('Zero-Secret Architecture', 'Nessun API key o password nel codice o nelle variabili d\'ambiente in produzione. Tutti i servizi Azure autenticati via Managed Identity con ruoli RBAC minimi: OpenAI User, AI Search Data Reader, Key Vault Secrets User.'),
        ('Private Endpoints', 'Tutti i servizi PaaS (OpenAI, AI Search, Key Vault, MongoDB) accessibili solo tramite Private Endpoints nella VNet. Nessuna esposizione su internet pubblico.'),
        ('Network Segmentation', 'VNet con subnet dedicate (api, data, private-endpoints) e NSGs con regole deny-all-inbound di default. Solo il traffico Container Apps ↔ servizi è consentito.'),
        ('Container Security', 'Dockerfile multi-stage con utenti non-root, health check integrati, immagini basate su Alpine/slim. Nginx con security headers (X-Frame-Options, X-Content-Type-Options, CSP).'),
        ('CAF Tagging', 'Tutti i moduli Bicep applicano tag obbligatori: Environment, Application, CostCenter, ManagedBy per governance e cost management.'),
        ('Audit Logging', 'Diagnostics settings su tutti i servizi con invio a Log Analytics (90 giorni retention). Application Insights per APM e distributed tracing.'),
    ]
    for title_text, desc in security_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{title_text}: ')
        run.bold = True
        p.add_run(desc)

    # ══════════════════════════════════════════════════════════════════════════
    # 9. DEVOPS E DEPLOYMENT
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('9. DevOps e Deployment', level=1)

    doc.add_heading('Local Development', level=2)
    doc.add_paragraph(
        'Per lo sviluppo locale è sufficiente Docker Compose che avvia 3 servizi: '
        'API (FastAPI con hot reload), Client (Nginx + React), MongoDB. '
        'In alternativa, la modalità Demo non richiede nessun backend — basta npm run dev nel client.'
    )

    doc.add_heading('Azure Deployment', level=2)
    doc.add_paragraph(
        'Lo script scripts/deploy.sh orchestra il deployment completo:'
    )
    deploy_steps = [
        'Login Azure e verifica subscription',
        'Deployment Bicep a scope subscription (tutti i moduli infrastrutturali)',
        'Login ACR, build e push immagini Docker (API + Client)',
        'Update Container App con la nuova immagine',
        'Output degli URL di accesso (API, docs, health)',
    ]
    for step in deploy_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph(
        'Lo script scripts/setup-rbac.sh completa la configurazione assegnando i ruoli RBAC '
        'alla Managed Identity del Container App: OpenAI User, AI Search Data Reader, Key Vault Secrets User.'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 10. SCENARI D'USO
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('10. Scenari d\'Uso Supportati', level=1)

    add_styled_table(doc,
        ['Scenario', 'Descrizione', 'Agenti Coinvolti'],
        [
            ['AWS → Azure Migration', 'Migrazione completa: service mapping, wave plan, cost comparison, strategy 6Rs', 'Tutti i 7 agenti'],
            ['GCP → Azure Migration', 'GKE→AKS, Pub/Sub→Event Hubs, Cloud SQL→Azure DB, IAM mapping', 'Tutti i 7 agenti'],
            ['On-Premises Modernization', 'Lift-and-shift vs re-architect, .NET upgrade, SQL migration, AD→Azure AD', 'Tutti i 7 agenti'],
            ['Cost Optimization', 'Right-sizing senza migrazione, reserved instances, orphaned resources cleanup', 'Code, Infra, Cost, Quality'],
            ['WAF Assessment', 'Scoring architettura su 5 pilastri con remediation plan prioritizzato', 'Code, Infra, WAF, Quality'],
            ['App Modernization', 'Monolith→microservices o serverless, containerization readiness', 'Code, Infra, Migration, GAP, Quality'],
            ['Code & IaC Quality Gate', 'Analisi statica SonarQube-level su tutti gli artefatti: bugs, vulnerabilities, smells, hotspots, debt', 'Code, Infra, Quality'],
        ],
        col_widths=[3.5, 7.5, 5.5]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 11. STACK TECNOLOGICO
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('11. Stack Tecnologico', level=1)

    add_styled_table(doc,
        ['Componente', 'Tecnologia', 'Versione', 'Motivazione'],
        [
            ['Frontend', 'React + Vite + Tailwind CSS', '18.3 / 6.0 / 3.4', 'SPA moderna, fast HMR, utility-first CSS, dark mode nativo'],
            ['Backend', 'FastAPI (Python)', '3.11', 'Native asyncio per agent parallelism, auto-docs OpenAPI, Pydantic validation'],
            ['LLM', 'Azure OpenAI GPT-4o', '2025-01-01', 'Best reasoning per analisi architetturale, supporto JSON response format'],
            ['Agent Framework', 'Azure AI Agent Service', '—', 'Persistence, threading, file search nativi. Integrazione Azure AI Foundry'],
            ['Knowledge Base', 'Azure AI Search', '—', 'Ricerca semantica per RAG su documenti CAF/WAF/migration patterns'],
            ['Database', 'MongoDB 7', '7.x', 'Schemaless, open-source, portabile. Su Container Apps con Azure Files (no license PaaS)'],
            ['Auth', 'Managed Identity (RBAC)', '—', 'Zero secrets in produzione. DefaultAzureCredential per dev locale'],
            ['Compute', 'Azure Container Apps', '—', 'Auto-scale 0-10, consumption plan, Dapr sidecar opzionale, managed TLS'],
            ['IaC', 'Bicep', '—', 'Formato nativo Azure, strong typing, modularizzazione, parametri per ambiente'],
            ['Containers', 'Docker (multi-stage)', '—', 'uv per Python deps, non-root users, health checks, Nginx per SPA'],
            ['Quality', 'SonarQube (integration)', '—', 'sonar-project.properties configurato + Quality Analyzer agent AI-powered'],
        ],
        col_widths=[2.5, 4, 2.5, 7.5]
    )

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('— Fine Documento —')
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(10)
    run.italic = True

    doc.add_paragraph()
    gen = doc.add_paragraph()
    gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = gen.add_run(
        'Documento generato automaticamente da Efesto AI Fabryc v1.0.0\n'
        'Powered by Azure AI Foundry · Follows CAF & WAF'
    )
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(9)

    # ── Save ──────────────────────────────────────────────────────────────────
    output_path = Path(__file__).parent.parent / 'docs' / 'Azure_Foundry_Architect_Framework_Documentation.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Documentation generated: {output_path}")
    return output_path


if __name__ == '__main__':
    generate_documentation()
